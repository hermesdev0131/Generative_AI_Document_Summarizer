import os
from summarizer import summarize_doc_sections
from template_sections import extract_template_sections, create_template_sections

from dotenv import load_dotenv
#NEW
from docx import Document
#NEW
from docx.shared import RGBColor, Pt, Inches
#NEW
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
# Load environment variables from .env file
load_dotenv(override=True)
# Set up OpenAI API key

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Model configuration
MODEL_TYPE = os.getenv("MODEL_TYPE", "openai").lower()  # Options: "openai", "azure", "llama"

# Set up model configuration based on the selected model type
model_config = {
    "model_type": MODEL_TYPE
}

if MODEL_TYPE == "openai":
    # OpenAI configuration
    model_config["api_key"] = os.getenv("OPENAI_API_KEY")
    if not model_config["api_key"]:
        raise ValueError("Please set the OPENAI_API_KEY environment variable.")
    
elif MODEL_TYPE == "azure":
    # Azure OpenAI configuration
    model_config["api_key"] = os.getenv("AZURE_OPENAI_API_KEY")
    model_config["endpoint"] = os.getenv("AZURE_OPENAI_ENDPOINT")
    model_config["deployment_name"] = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
    
    if not all([model_config["api_key"], model_config["endpoint"], model_config["deployment_name"]]):
        raise ValueError("Please set all required Azure OpenAI environment variables: "
                         "AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_DEPLOYMENT_NAME")
    
elif MODEL_TYPE == "llama":
    # Llama configuration
    model_config["model_path"] = os.getenv("LLAMA_MODEL_PATH")
    
    if not model_config["model_path"]:
        raise ValueError("Please set the LLAMA_MODEL_PATH environment variable.")
    
    if not os.path.exists(model_config["model_path"]):
        raise FileNotFoundError(f"Llama model not found at '{model_config['model_path']}'")
else:
    raise ValueError(f"Unsupported model type: {MODEL_TYPE}. Choose from 'openai', 'azure', or 'llama'.")

current_dir = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DOC_PATH = os.path.join(current_dir, "template_doc", "template.docx")
INPUT_DIR = "input_doc"
TEMPLATE_PATH = "template_doc/1759_Prospect Curator_Template.pdf"
OUTPUT_DIR = "output_doc"
if not os.path.exists(TEMPLATE_DOC_PATH):
    raise FileNotFoundError(f"Template file not found at '{TEMPLATE_DOC_PATH}'")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# NEW - This code allows us to either use PDF template or provide plaintext guidelines
USE_PDF_TEMPLATE = False  # Set to False to use plaintext templates

if USE_PDF_TEMPLATE:
    # Get summarization instructions from template PDF
    instructions = extract_template_sections(TEMPLATE_PATH)
else:
    # Example of using plaintext templates
    instructions = create_template_sections(TEMPLATE_DOC_PATH)

SECTION_ORDER = ["Prospect Curator", "Back of the Napkin", "The Pitch", "Rewards", "Client Needs", "The Elixir", "Insurance+", "Landmines", "Path to Close", "What Will Clients Ask?", "Sustain Success", "Look Ahead"]

def create_formatted_docx(summaries, output_path):
    doc = Document()
    
    # Set default paragraph spacing and font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(12)
    
    # Define section header style
    def add_section_header(text):
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue color
        header.space_before = Pt(24)
        header.space_after = Pt(12)
        
    # Add each section
    for section in SECTION_ORDER:
        if section in summaries:
            add_section_header(section)
            
            # Add content with proper paragraph formatting
            paragraphs = summaries[section].split('\n\n')
            for i, p in enumerate(paragraphs):
                if p.strip():
                    content = doc.add_paragraph()
                    # Justify all paragraphs except the last one
                    if i < len(paragraphs) - 1:
                        content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        content.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    content.add_run(p.strip())
    
    # Save the document
    doc.save(output_path)


for filename in os.listdir(INPUT_DIR):
    if filename.lower().endswith((".pdf", ".docx")):
        input_path = os.path.join(INPUT_DIR, filename)
        base_name = os.path.splitext(filename)[0]
        output_path = os.path.join(OUTPUT_DIR, f"{base_name}_summary.docx")

        # print(f"Processing {filename}...")
        logging.info(f"Using model type: {MODEL_TYPE}")

        # Get summaries for all sections
        summaries = summarize_doc_sections(input_path, model_config, instructions)

        # NEW - Create formatted Word document
        create_formatted_docx(summaries, output_path)

        # print(f"Created: {output_path}")
        logging.info(f"Successfully created summary at {output_path}")